"""
app/components/resume_files.py
Phase 2 file I/O for the Resume Workshop.

Two responsibilities:
  1. Extract plain text from an uploaded resume (PDF or DOCX) so Claude can read it.
  2. Export tailored markdown output back to a formatted .docx for download.

Kept dependency-light: pypdf for PDFs, python-docx for Word in/out.
"""

import io
import re


# ── Text extraction (upload → text) ───────────────────────────────────────────

def extract_text(filename: str, data: bytes) -> str:
    """
    Extract plain text from an uploaded resume file.

    Supports .pdf and .docx (and .txt/.md as a passthrough).
    Raises ValueError on an unsupported extension, and RuntimeError if the
    file is recognized but can't be parsed (corrupt, image-only PDF, etc.).
    """
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        text = _extract_pdf(data)
    elif name.endswith(".docx"):
        text = _extract_docx(data)
    elif name.endswith((".txt", ".md")):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError(
            "Unsupported file type. Upload a .pdf, .docx, .txt, or .md resume."
        )

    text = _tidy(text)
    if not text.strip():
        raise RuntimeError(
            "No readable text found. If this is a scanned or image-only PDF, "
            "export a text-based version and try again."
        )
    return text


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs]

    # Pull text out of tables too — many resumes lay out contact info in a table.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append("  ".join(cells))

    return "\n".join(lines)


def _tidy(text: str) -> str:
    """Normalize whitespace: collapse runs of blank lines, strip trailing spaces."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── DOCX export (markdown → .docx bytes) ───────────────────────────────────────

def markdown_to_docx(markdown_text: str, title: str = "") -> bytes:
    """
    Render a lightweight subset of markdown to a .docx and return the bytes.

    Handles: #/##/### headings, - or * bullets, blank-line paragraph breaks,
    and inline **bold**. Anything fancier is written as a plain paragraph — the
    goal is a clean, ATS-friendly Word document, not a full markdown renderer.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Base font: keep it simple and ATS-safe.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for raw_line in (markdown_text or "").split("\n"):
        line = raw_line.rstrip()

        if not line.strip():
            continue  # skip blank lines; Word spacing handles separation

        # Horizontal rules → skip (no clean docx equivalent worth the noise)
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line.strip()):
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            doc.add_heading(heading.group(2).strip(), level=level)
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            _add_runs(para, bullet.group(1).strip())
            continue

        para = doc.add_paragraph()
        _add_runs(para, line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, honoring **bold** spans."""
    # Split on **bold** markers, keeping the delimiters' content.
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        # Odd-index chunks were captured inside ** ** → bold.
        if i % 2 == 1:
            run.bold = True
